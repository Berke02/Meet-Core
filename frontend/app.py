import os
import re
import requests
import streamlit as st
import pandas as pd
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

# =========================================================
# PAGE CONFIG
# =========================================================

st.set_page_config(page_title="Meet Core", page_icon="🎙️", layout="wide")


# =========================================================
# API CONFIG
# =========================================================

DEFAULT_API_BASE_URL = os.getenv("MEET_CORE_API_URL", "http://127.0.0.1:8000")


def normalize_api_url(url: str) -> str:
    return url.strip().rstrip("/")


# =========================================================
# SPEAKER DISPLAY HELPERS
# =========================================================

SPEAKER_PATTERN = re.compile(r"\bSPEAKER_(\d+)\b")


def extract_speaker_labels(text: str | None) -> list[str]:
    if not text:
        return []

    labels = []
    seen = set()

    for match in SPEAKER_PATTERN.finditer(text):
        label = f"SPEAKER_{int(match.group(1)):02d}"

        if label not in seen:
            labels.append(label)
            seen.add(label)

    return labels


def build_display_speaker_map(
    transcript_text: str | None,
    speaker_aliases: dict[str, str] | None = None,
) -> dict[str, str]:
    """
    Speaker etiketlerini kullanıcıya gösterilecek isimlere çevirir.

    Örnek:
    speaker_aliases = {"SPEAKER_01": "Merve"}

    Transcript içinde SPEAKER_01, SPEAKER_02, SPEAKER_00 varsa:
    SPEAKER_01 -> Merve
    SPEAKER_02 -> UNKNOWN_01
    SPEAKER_00 -> UNKNOWN_02

    UNKNOWN numarası SPEAKER numarasına göre değil,
    çözülemeyen konuşmacı sırasına göre verilir.
    """

    speaker_aliases = speaker_aliases or {}
    speaker_labels = extract_speaker_labels(transcript_text)

    display_map = {}
    unknown_counter = 1

    for speaker_label in speaker_labels:
        alias = speaker_aliases.get(speaker_label, "").strip()

        if alias:
            display_map[speaker_label] = alias
        else:
            display_map[speaker_label] = f"UNKNOWN_{unknown_counter:02d}"
            unknown_counter += 1

    return display_map


def replace_speaker_labels(value: str | None, display_map: dict[str, str]) -> str | None:
    if value is None:
        return None

    text = str(value)

    for raw_label, display_label in display_map.items():
        text = text.replace(raw_label, display_label)

    return text


def render_speaker_mapping_controls(transcript_text: str | None) -> dict[str, str]:
    speaker_labels = extract_speaker_labels(transcript_text)

    if "speaker_aliases" not in st.session_state:
        st.session_state.speaker_aliases = {}

    display_map = build_display_speaker_map(
        transcript_text=transcript_text,
        speaker_aliases=st.session_state.speaker_aliases,
    )

    if not speaker_labels:
        return display_map

    with st.expander("👥 Konuşmacı isimlerini düzenle", expanded=False):
        st.caption(
            "Gerçek ismini bildiğiniz konuşmacıları yazın. "
            "Boş bıraktıklarınız UNKNOWN_01, UNKNOWN_02 şeklinde gösterilir."
        )

        updated_aliases = {}

        for speaker_label in speaker_labels:
            current_value = st.session_state.speaker_aliases.get(speaker_label, "")

            updated_aliases[speaker_label] = st.text_input(
                label=f"{speaker_label} için görünen isim",
                value=current_value,
                placeholder=display_map.get(speaker_label, ""),
                key=f"speaker_alias_input_{speaker_label}",
            )

        st.session_state.speaker_aliases = updated_aliases

    return build_display_speaker_map(
        transcript_text=transcript_text,
        speaker_aliases=st.session_state.speaker_aliases,
    )

# =========================================================
# DOCUMENT EXTRACTION
# =========================================================

def extract_document_text(uploaded_file) -> str:
    """PDF, DOCX veya TXT dosyasından metin çıkarır."""

    if uploaded_file is None:
        raise ValueError("Doküman dosyası bulunamadı.")

    filename = uploaded_file.name or "document"
    suffix = Path(filename).suffix.lower()
    file_bytes = uploaded_file.getvalue()

    if not file_bytes:
        raise ValueError("Yüklenen doküman boş.")

    if suffix == ".txt":
        text = _decode_text_file(file_bytes)

    elif suffix == ".pdf":
        text = _extract_pdf_text(file_bytes)

    elif suffix == ".docx":
        text = _extract_docx_text(file_bytes)

    else:
        raise ValueError(
            f"Desteklenmeyen doküman formatı: {suffix}. "
            "Desteklenen formatlar: PDF, DOCX ve TXT."
        )

    normalized_text = text.strip()

    if not normalized_text:
        raise ValueError(
            "Dokümandan metin çıkarılamadı. "
            "PDF taranmış görüntülerden oluşuyorsa OCR gerekir."
        )

    return normalized_text


def _decode_text_file(file_bytes: bytes) -> str:
    """TXT dosyasını yaygın Türkçe karakter kodlamalarıyla çözer."""

    encodings = (
        "utf-8-sig",
        "utf-8",
        "cp1254",
        "latin-1",
    )

    for encoding in encodings:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise ValueError("TXT dosyasının karakter kodlaması çözülemedi.")


def _extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    page_texts: list[str] = []

    for page_number, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:
            raise ValueError(
                f"PDF'in {page_number}. sayfası okunamadı: {exc}"
            ) from exc

        if page_text.strip():
            page_texts.append(page_text.strip())

    return "\n\n".join(page_texts)


def _extract_docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            text_parts.append(paragraph_text)

    # Word tablolarındaki metinleri de dahil et.
    for table in document.tables:
        for row in table.rows:
            cell_values = [
                cell.text.strip()
                for cell in row.cells
            ]

            if any(cell_values):
                text_parts.append(" | ".join(cell_values))

    return "\n".join(text_parts)

# =========================================================
# BACKEND CALLS
# =========================================================
def call_text_analysis(
    api_base_url: str,
    meeting_text: str,
    participant_count: int,
) -> dict:
    response = requests.post(
        f"{api_base_url}/api/meetings/analyze",
        json={
            "meeting_text": meeting_text,
            "participant_count": participant_count,
        },
        timeout=180,
    )

    if not response.ok:
        raise RuntimeError(
            f"Backend error {response.status_code}: {response.text}"
        )

    return response.json()


def call_audio_analysis(
    api_base_url: str,
    uploaded_file,
    participant_count: int,
) -> dict:
    files = {
        "file": (
            uploaded_file.name,
            uploaded_file.getvalue(),
            uploaded_file.type or "application/octet-stream",
        )
    }

    form_data = {
        "participant_count": str(participant_count),
    }

    response = requests.post(
        f"{api_base_url}/api/meetings/analyze-audio",
        files=files,
        data=form_data,
        timeout=600,
    )

    if not response.ok:
        raise RuntimeError(
            f"Backend error {response.status_code}: {response.text}"
        )

    return response.json()


# =========================================================
# DATA HELPERS
# =========================================================

def get_action_items(analysis: dict) -> list[dict]:
    return analysis.get("action_items", []) or []


def is_high_risk_action(item: dict) -> bool:
    metrics = item.get("metrics") or {}

    risk_level = item.get("risk_level")
    urgency_score = float(metrics.get("urgency_score") or 0)
    ambiguity_score = float(metrics.get("ambiguity_score") or 0)

    return (
        risk_level == "high"
        or urgency_score >= 70
        or ambiguity_score >= 50
    )


def normalize_display_name(name: str | None, display_map: dict[str, str]) -> str | None:
    if name is None:
        return None

    value = replace_speaker_labels(str(name), display_map)

    if value is None:
        return None

    value = value.strip()

    if not value or value in {"-", "None", "null", "NULL"}:
        return None

    return value


def build_participant_dataframe(
    participants: list[dict],
    action_items: list[dict],
    decisions: list[dict],
    open_questions: list[dict],
    transcript_text: str | None,
    display_map: dict[str, str],
) -> pd.DataFrame:
    rows_by_name: dict[str, dict] = {}

    def add_person(name: str | None, role: str | None = None) -> None:
        display_name = normalize_display_name(name, display_map)

        if not display_name:
            return

        if display_name not in rows_by_name:
            rows_by_name[display_name] = {
                "name": display_name,
                "role": role,
            }
        elif not rows_by_name[display_name].get("role") and role:
            rows_by_name[display_name]["role"] = role

    # Canonical kaynak her zaman backend participants olmalı.
    for participant in participants or []:
        add_person(
            participant.get("name"),
            participant.get("role"),
        )

    # Eski backend veya metin analizleri için fallback.
    if not rows_by_name:
        for item in action_items or []:
            add_person(item.get("owner"))

        for decision in decisions or []:
            add_person(decision.get("owner"))

        for question in open_questions or []:
            add_person(question.get("owner"))

    return pd.DataFrame(list(rows_by_name.values()))

def build_action_dataframe(
    action_items: list[dict],
    display_map: dict[str, str] | None = None,
) -> pd.DataFrame:
    display_map = display_map or {}
    rows = []

    for item in action_items:
        metrics = item.get("metrics") or {}

        rows.append(
            {
                "Sorumlu": replace_speaker_labels(item.get("owner"), display_map) or "-",
                "Görev": item.get("task") or "-",
                "Deadline": item.get("due_date") or "-",
                "Öncelik": item.get("priority") or "-",
                "Risk Seviyesi": item.get("risk_level") or "-",
                "Muğlaklık Skoru": metrics.get("ambiguity_score"),
                "Aciliyet Skoru": metrics.get("urgency_score"),
                "Aksiyon Yoğunluğu": metrics.get("action_density_score"),
                "Kaynak Cümle": replace_speaker_labels(item.get("source_sentence"), display_map) or "-",
            }
        )

    return pd.DataFrame(rows)

def build_grouped_action_dataframe(action_df: pd.DataFrame) -> pd.DataFrame:
    """
    Aynı sorumluya ait görevleri tek satırda toplar.
    Özellikle UNKNOWN_01 gibi aynı bilinmeyen konuşmacıya ait görevlerin
    ayrı ayrı satırlarda görünmesini engeller.
    """

    if action_df.empty:
        return action_df

    grouped_rows = []

    for owner, group in action_df.groupby("Sorumlu", dropna=False):
        tasks = "\n".join(
            f"- {task}" for task in group["Görev"].dropna().astype(str).tolist()
        )

        source_sentences = "\n".join(
            f"- {source}" for source in group["Kaynak Cümle"].dropna().astype(str).tolist()
        )

        deadlines = sorted(
            {
                str(value)
                for value in group["Deadline"].dropna().tolist()
                if str(value).strip() not in {"", "-"}
            }
        )

        priorities = sorted(
            {
                str(value)
                for value in group["Öncelik"].dropna().tolist()
                if str(value).strip() not in {"", "-"}
            }
        )

        risk_levels = sorted(
            {
                str(value)
                for value in group["Risk Seviyesi"].dropna().tolist()
                if str(value).strip() not in {"", "-"}
            }
        )

        grouped_rows.append(
            {
                "Sorumlu": owner,
                "Görev Sayısı": len(group),
                "Görevler": tasks,
                "Deadline": ", ".join(deadlines) if deadlines else "-",
                "Öncelik": ", ".join(priorities) if priorities else "-",
                "Risk Seviyesi": ", ".join(risk_levels) if risk_levels else "-",
                "Ortalama Muğlaklık": round(group["Muğlaklık Skoru"].fillna(0).mean(), 2),
                "Ortalama Aciliyet": round(group["Aciliyet Skoru"].fillna(0).mean(), 2),
                "Ortalama Aksiyon Yoğunluğu": round(group["Aksiyon Yoğunluğu"].fillna(0).mean(), 2),
                "Kaynak Cümleler": source_sentences,
            }
        )

    return pd.DataFrame(grouped_rows)

def build_decision_dataframe(
    decisions: list[dict],
    display_map: dict[str, str] | None = None,
) -> pd.DataFrame:
    display_map = display_map or {}
    rows = []

    for decision in decisions or []:
        rows.append(
            {
                "Karar": decision.get("decision") or "-",
                "Sorumlu": replace_speaker_labels(decision.get("owner"), display_map) or "-",
                "Güven": decision.get("confidence"),
                "Kaynak Cümle": replace_speaker_labels(decision.get("source_sentence"), display_map) or "-",
            }
        )

    return pd.DataFrame(rows)


def build_open_questions_dataframe(
    open_questions: list[dict],
    display_map: dict[str, str] | None = None,
) -> pd.DataFrame:
    display_map = display_map or {}
    rows = []

    for question in open_questions or []:
        rows.append(
            {
                "Soru": question.get("question") or "-",
                "Sorumlu": replace_speaker_labels(question.get("owner"), display_map) or "-",
            }
        )

    return pd.DataFrame(rows)


# =========================================================
# RESULT RENDERING
# =========================================================

def render_analysis_result(analysis: dict, transcript_text: str | None = None) -> None:
    transcript_display_map = render_speaker_mapping_controls(
        transcript_text
    )

    # Structured backend çıktısında backend isimleri kanoniktir.
    analysis_display_map: dict[str, str] = {}

    action_items = get_action_items(analysis)
    decisions = analysis.get("decisions", []) or []
    participants = analysis.get("participants", []) or []
    open_questions = analysis.get("open_questions", []) or []

    participant_df = build_participant_dataframe(
        participants=participants,
        action_items=action_items,
        decisions=decisions,
        open_questions=open_questions,
        transcript_text=transcript_text,
        display_map=analysis_display_map,
    )

    high_risk_count = sum(1 for item in action_items if is_high_risk_action(item))

    ambiguity_scores = [
        float((item.get("metrics") or {}).get("ambiguity_score") or 0)
        for item in action_items
    ]

    average_clarity = 100 - (sum(ambiguity_scores) / len(ambiguity_scores)) if ambiguity_scores else 100
    average_clarity = round(max(0, min(100, average_clarity)), 1)

    st.success("✅ Toplantı başarıyla analiz edildi!")

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Katılımcı", f"{len(participant_df)} Kişi")
    m2.metric("Çıkarılan Görev", f"{len(action_items)} Adet")
    m3.metric("Yüksek Riskli Görev", f"{high_risk_count} Adet")
    m4.metric("Görev Netliği", f"%{average_clarity}")

    st.divider()

    tab1, tab2, tab3, tab4, tab5 = st.tabs(
        [
            "📝 Özet & Kararlar",
            "🎯 Görevler ve Risk Skorları",
            "❓ Açık Sorular",
            "🎧 Transkript",
            "🧾 Backend Ham Çıktı",
        ]
    )

    with tab1:
        st.subheader("Toplantı Özeti")
        summary = replace_speaker_labels(analysis.get("summary"), analysis_display_map)
        st.info(summary or "Özet bulunamadı.")

        st.subheader("Ana Konular")
        key_topics = analysis.get("key_topics", []) or []

        if key_topics:
            for topic in key_topics:
                st.write(f"- {replace_speaker_labels(topic, analysis_display_map)}")
        else:
            st.write("Ana konu bulunamadı.")

        st.subheader("Katılımcılar")

        if not participant_df.empty:
            st.dataframe(participant_df, use_container_width=True)
        else:
            st.write("Katılımcı bulunamadı.")

        st.subheader("Kararlar")
        decision_df = build_decision_dataframe(decisions, analysis_display_map)

        if not decision_df.empty:
            st.dataframe(decision_df, use_container_width=True)
        else:
            st.warning(
                "Bu toplantıdan açık bir karar çıkarılmadı. "
                "Metin daha çok aksiyon/görev cümleleri içeriyor olabilir."
            )

    with tab2:
        st.subheader("Aksiyon Maddeleri ve Metrikler")

        action_df = build_action_dataframe(action_items, analysis_display_map)

        if not action_df.empty:
            grouped_action_df = build_grouped_action_dataframe(action_df)
            st.dataframe(grouped_action_df, use_container_width=True)
            st.subheader("Aciliyet Skoru Dağılımı")
            chart_df = action_df[["Görev", "Aciliyet Skoru"]].copy()
            chart_df = chart_df.dropna()

            if not chart_df.empty:
                st.bar_chart(chart_df.set_index("Görev"))
        else:
            st.write("Aksiyon maddesi bulunamadı.")

    with tab3:
        st.subheader("Açık Sorular")

        question_df = build_open_questions_dataframe(open_questions, analysis_display_map)

        if not question_df.empty:
            st.dataframe(question_df, use_container_width=True)
        else:
            st.write("Açık soru bulunamadı.")

    with tab4:
        st.subheader("Transkript")

        if transcript_text:
            display_transcript = replace_speaker_labels(transcript_text, transcript_display_map)
            st.text_area("Ses kaydından çıkarılan transkript", display_transcript, height=300)
        else:
            st.write("Bu analiz metin girişi üzerinden yapıldı; ayrı transkript çıktısı bulunmuyor.")

    with tab5:
        st.subheader("Backend Ham JSON Çıktısı")
        st.json(analysis)


# =========================================================
# SIDEBAR
# =========================================================

with st.sidebar:
    st.header("📌 Menü")

    secilen_sayfa = st.radio(
        "Sayfa Seçin:",
        ["Yeni Toplantı Analizi", "Risk Dashboard", "Geçmiş Toplantılar"],
    )

    st.divider()

    api_base_url = st.text_input(
        "Backend API URL",
        value=DEFAULT_API_BASE_URL,
    )
    api_base_url = normalize_api_url(api_base_url)

    st.info("🤖 Agent Durumu: Hazır")


# =========================================================
# SESSION STATE
# =========================================================

if "last_analysis" not in st.session_state:
    st.session_state.last_analysis = None

if "last_transcript" not in st.session_state:
    st.session_state.last_transcript = None

if "speaker_aliases" not in st.session_state:
    st.session_state.speaker_aliases = {}


# =========================================================
# PAGE 1: NEW MEETING ANALYSIS
# =========================================================

if secilen_sayfa == "Yeni Toplantı Analizi":
    st.title("🎙️ Meet Core: Yeni Toplantı Analizi")
    st.markdown(
        "Toplantı ses kaydını yükleyin veya notları yapıştırın. "
        "Yapay zeka ajanları sizin için aksiyonları, kararları ve risk metriklerini çıkarsın."
    )
    st.divider()

    col1, col2 = st.columns(2)

    with col1:
        uploaded_file = st.file_uploader(
            "🎵 Ses Kaydı Yükle",
            type=["mp3", "wav", "m4a", "mp4", "webm"],
            key="audio_upload",
        )

        document_file = st.file_uploader(
            "📄 Doküman Yükle",
            type=["pdf", "docx", "txt"],
            key="document_upload",
            help=(
                "Toplantı notlarını PDF, Word (.docx) veya TXT "
                "formatında yükleyebilirsiniz."
            ),
        )

    with col2:
        transcript = st.text_area(
            "📝 Veya Transkript Yapıştır",
            height=160,
            placeholder="Örn: Merve: LLM entegrasyonunu bugün tamamlayacağım...",
        )

    participant_count = st.number_input(
        "👥 Toplantıdaki toplam katılımcı sayısı",
        min_value=1,
        max_value=50,
        value=5,
        step=1,
        help="Adı tespit edilemeyen kişiler UNKNOWN olarak gösterilir.",
    )

    st.markdown("<br>", unsafe_allow_html=True)

    if st.button("🚀 Yapay Zeka ile Analiz Et", use_container_width=True):
        has_audio = uploaded_file is not None
        has_document = document_file is not None
        has_transcript = bool(transcript.strip())

        selected_source_count = sum(
            [
                has_audio,
                has_document,
                has_transcript,
            ]
        )

        if selected_source_count == 0:
            st.error(
                "Lütfen ses kaydı, doküman veya transkript kaynaklarından "
                "birini seçin."
            )

        elif selected_source_count > 1:
            st.error(
                "Aynı anda yalnızca bir analiz kaynağı kullanabilirsiniz. "
                "Ses kaydı, doküman veya transkriptten birini seçin."
            )

        else:
            try:
                with st.spinner(
                    "Agent'lar çalışıyor... "
                    "Analiz ve risk metrikleri hesaplanıyor..."
                ):
                    if uploaded_file:
                        backend_response = call_audio_analysis(
                            api_base_url=api_base_url,
                            uploaded_file=uploaded_file,
                            participant_count=int(participant_count),
                        )

                        analysis = backend_response.get("analysis")
                        transcript_text = backend_response.get("transcript_text")

                        if not analysis:
                            raise RuntimeError(
                                "Backend audio response içinde analysis alanı bulunamadı."
                            )

                    elif document_file:
                        document_text = extract_document_text(document_file)

                        analysis = call_text_analysis(
                            api_base_url=api_base_url,
                            meeting_text=document_text,
                            participant_count=int(participant_count),
                        )

                        transcript_text = document_text

                    else:
                        transcript_text = transcript.strip()

                        analysis = call_text_analysis(
                            api_base_url=api_base_url,
                            meeting_text=transcript_text,
                            participant_count=int(participant_count),
                        )

                    actual_participant_count = len(
                        analysis.get("participants", [])
                    )

                    if actual_participant_count != int(participant_count):
                        raise RuntimeError(
                            "Backend katılımcı sayısı hatalı. "
                            f"Beklenen={participant_count}, "
                            f"Gelen={actual_participant_count}"
                        )

                    st.session_state.last_analysis = analysis
                    st.session_state.last_transcript = transcript_text
                    st.session_state.speaker_aliases = {}

                render_analysis_result(
                    st.session_state.last_analysis,
                    st.session_state.last_transcript,
                )

            except Exception as exc:
                st.error("Analiz sırasında hata oluştu.")
                st.code(str(exc))

    elif st.session_state.last_analysis:
        render_analysis_result(
            st.session_state.last_analysis,
            st.session_state.last_transcript,
        )


# =========================================================
# PAGE 2: RISK DASHBOARD
# =========================================================

elif secilen_sayfa == "Risk Dashboard":
    st.title("📊 Risk ve Takip Dashboard'u")
    st.markdown("Son analiz edilen toplantıdaki görevlerin kişi bazlı risk ve metrik görünümü.")
    st.divider()

    if not st.session_state.last_analysis:
        st.warning("Henüz analiz edilmiş toplantı yok. Önce 'Yeni Toplantı Analizi' sayfasında bir analiz çalıştırın.")
    else:
        action_items = get_action_items(
            st.session_state.last_analysis
        )

        action_df = build_action_dataframe(
            action_items=action_items,
            display_map={},
        )

        if action_df.empty:
            st.info("Son analizde aksiyon maddesi bulunamadı.")
        else:
            owner_summary = (
                action_df.assign(
                    Yuksek_Risk=action_df.apply(
                        lambda row: (
                            row["Risk Seviyesi"] == "high"
                            or float(row["Aciliyet Skoru"] or 0) >= 70
                            or float(row["Muğlaklık Skoru"] or 0) >= 50
                        ),
                        axis=1,
                    )
                )
                .groupby("Sorumlu")
                .agg(
                    Açık_Görev_Sayısı=("Görev", "count"),
                    Yüksek_Riskli_Görev=("Yuksek_Risk", "sum"),
                    Ortalama_Aciliyet=("Aciliyet Skoru", "mean"),
                    Ortalama_Muğlaklık=("Muğlaklık Skoru", "mean"),
                )
                .reset_index()
            )

            st.subheader("👤 Kişi Bazlı Görev Özeti")
            st.dataframe(owner_summary, use_container_width=True)

            col_grafik1, col_grafik2 = st.columns(2)

            with col_grafik1:
                st.subheader("Toplam Görev Dağılımı")
                st.bar_chart(owner_summary.set_index("Sorumlu")["Açık_Görev_Sayısı"])

            with col_grafik2:
                st.subheader("Yüksek Riskli Görevler")
                st.bar_chart(owner_summary.set_index("Sorumlu")["Yüksek_Riskli_Görev"])

            st.subheader("Detaylı Görev Listesi")
            grouped_action_df = build_grouped_action_dataframe(action_df)
            st.dataframe(grouped_action_df, use_container_width=True)


# =========================================================
# PAGE 3: MEETING HISTORY
# =========================================================

elif secilen_sayfa == "Geçmiş Toplantılar":
    st.title("🗂️ Geçmiş Toplantılar Arşivi")
    st.markdown(
        "Şimdilik son analiz edilen toplantı session üzerinde tutuluyor. "
        "Kalıcı kayıt için backend tarafında database entegrasyonu gerekir."
    )
    st.divider()

    if not st.session_state.last_analysis:
        st.warning("Henüz görüntülenecek toplantı analizi yok.")
    else:
        st.subheader("Son Analiz Edilen Toplantı")
        render_analysis_result(
            st.session_state.last_analysis,
            st.session_state.last_transcript,
        )